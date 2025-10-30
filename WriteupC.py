from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path

# === Theme constants ===
KU_CRIMSON = RGBColor(0xA6, 0x19, 0x2E)  # Korea University crimson
BODY_FONT = "Times New Roman"

def set_doc_defaults(doc: Document):
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # Heading styles
    for name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13), ("Heading 4", 12)]:
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = KU_CRIMSON
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.space_before = Pt(6 if name != "Heading 1" else 0)
        st.paragraph_format.space_after = Pt(4)

    # Create "Bibliography" style for hanging indent
    if "Bibliography" not in [s.name for s in doc.styles]:
        bibl = doc.styles.add_style("Bibliography", 1)  # paragraph style
    else:
        bibl = doc.styles["Bibliography"]
    bibl.font.name = BODY_FONT
    bibl.font.size = Pt(11)
    bibl.paragraph_format.line_spacing = 1.0
    bibl.paragraph_format.space_before = Pt(0)
    bibl.paragraph_format.space_after = Pt(6)
    bibl.paragraph_format.left_indent = Inches(0.5)
    bibl.paragraph_format.first_line_indent = Inches(-0.5)

def set_cell_borders(tc, color_hex="A6192E"):
    tc_pr = tc._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        e = borders.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            borders.append(e)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color_hex)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    r = p.runs[0]
    r.bold = bold
    r.italic = italic
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

def add_quote_box(doc, text):
    """Render a quoted block inside a single-cell table box in KU color."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.italic = True
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

# Build document
doc = Document()
set_doc_defaults(doc)

# Title
title = add_heading(doc, "SK Telecom Data Breach Analysis: Security Principle Violations and Design Recommendations", 1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Korea University — COSE354 Secure Coding Practice")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(6)

# Section 1
add_heading(doc, "1) Violation of Security Principles", 2)

# Plain-text Credential Storage
add_heading(doc, "Plain-text Credential Storage", 3)
add_para(doc, 
    'Violation Details: The investigation revealed systematic plain-text credential storage, where '
    '"Server A held—in unencrypted form—the IDs and passwords of other management-subnet hosts" and '
    '"Server B, in turn, stored plain-text admin credentials for the HSS management server" (MSIT, 2025). '
    'This fundamental failure allowed attackers to easily harvest and reuse credentials across multiple systems, '
    'enabling the breach to escalate from initial access to critical authentication systems.'
)
add_para(doc, "Principle Violations:", bold=True)
add_bullets(doc, [
    "Safe Defaults: Systems defaulted to storing credentials in plain text rather than encrypted form as the baseline security posture",
    "Complete Mediation: Credential validation mechanisms failed to enforce proper protection throughout the credential lifecycle, allowing stolen credentials to be reused across system boundaries",
])
add_para(doc,
    "CWE Reference: According to MITRE CWE-256/312, plain-text storage of credentials violates fundamental security practices by exposing sensitive authentication data to unauthorized access, creating a single point of failure that can compromise entire systems."
)

# Failure to Encrypt Critical Data
add_heading(doc, "Failure to Encrypt Critical Data", 3)
add_para(doc,
    'Violation Details: The breach exposed that "the USIM authentication key (Ki) was stored in plain text, contrary to GSMA recommendations and the practice of peers" (MSIT, 2025). '
    "This affected 9.82 GB of USIM data covering approximately 26.96 million records, creating what one analysis called "
    '"a wake-up call for digital security and corporate responsibility" (Law and Ethics in Tech, 2025).'
)
add_para(doc, "Principle Violations:", bold=True)
add_bullets(doc, [
    "Safe Defaults: Systems defaulted to unencrypted storage for highly sensitive authentication data rather than implementing encryption by default",
    "Open Design: Security relied on secrecy of data storage locations rather than proper cryptographic protection of the data itself",
])

# Inadequate Security Governance and Access Control
add_heading(doc, "Inadequate Security Governance and Access Control", 3)
add_para(doc,
    'Violation Details: The investigation found "fragmented governance" where "SK Telecom\'s CISO covered only the IT domain (57% of assets), '
    'leaving the network domain (43%) under separate supervision" (MSIT, 2025). This violated the Network Act requirement for unified security oversight '
    "and created security gaps that attackers exploited."
)
add_para(doc, "Principle Violations:", bold=True)
add_bullets(doc, [
    "Separation of Duties: Security responsibilities were improperly divided without clear accountability structures",
    "Least Privilege: The attacker maintained \"long-term persistence\" because \"administrative passwords had no expiry and had not been rotated for years,\" demonstrating excessive privilege duration",
])

# Incomplete Security Monitoring and Response
add_heading(doc, "Incomplete Security Monitoring and Response", 3)
add_para(doc,
    'Violation Details: During a 2022 incident response, SKT "reviewed only one of six available log files, missing the attacker\'s access traces" (MSIT, 2025). '
    "This incomplete investigation allowed the breach to continue undetected for years, with the company also failing to meet the 24-hour statutory reporting requirement."
)
add_para(doc, "Principle Violations:", bold=True)
add_bullets(doc, [
    "Complete Mediation: Security monitoring did not comprehensively cover all access attempts and system activities",
    "Ease of Use: Complex log analysis procedures led to incomplete investigations of critical security events, violating the principle that security mechanisms should not hinder proper oversight",
])

# Section 2
add_heading(doc, "2) Design Suggestions Based on Security Principles", 2)

# Comprehensive Credential Management System
add_heading(doc, "Comprehensive Credential Management System", 3)
add_para(doc, "Safe Defaults Implementation:", bold=True)
add_bullets(doc, [
    "Implement mandatory encryption for all stored credentials using strong cryptographic algorithms as the default configuration",
    "Deploy automated credential rotation with maximum lifespan policies, eliminating long-lived credentials that enabled \"long-term persistence\" in the attack",
    "As MSIT recommended, \"restrict any recording of passwords and, if unavoidable, store them in encrypted form while introducing multi-factor authentication\" (MSIT, 2025)",
])
add_para(doc,
    "CWE Mitigation: Following CWE-798 recommendations, eliminate hard-coded credentials and implement secure credential management systems with proper key rotation to prevent credential reuse attacks."
)

# Data Protection by Design
add_heading(doc, "Data Protection by Design", 3)
add_para(doc, "Safe Defaults and Open Design:", bold=True)
add_bullets(doc, [
    "Encrypt all sensitive data including USIM authentication keys using industry-standard algorithms as the default configuration",
    "Implement proper key management with regular rotation schedules and hardware security modules for master key protection",
    "As identified in the report, \"encrypt Ki and other key fields in line with domestic law and GSMA guidance\" (MSIT, 2025), ensuring safe defaults are enforced",
])
add_para(doc,
    "CWE Guidance: According to CWE-522, implement sufficiently protected credentials through proper encryption and access controls, preventing the plain-text storage that enabled this breach."
)

# Unified Security Governance Framework
add_heading(doc, "Unified Security Governance Framework", 3)
add_para(doc, "Separation of Duties and Least Privilege:", bold=True)
add_bullets(doc, [
    'Establish single CISO with enterprise-wide authority and direct CEO reporting as required by "Article 45-3 of the Network Act" (MSIT, 2025)',
    "Implement role-based access control with minimum necessary privileges across all domains, preventing credential reuse across system boundaries",
    "Deploy zero-trust architecture with continuous authentication validation to enforce complete mediation of all access attempts",
])

# Comprehensive Security Monitoring
add_heading(doc, "Comprehensive Security Monitoring", 3)
add_para(doc, "Complete Mediation and Ease of Use:", bold=True)
add_bullets(doc, [
    "Implement centralized security monitoring with automated threat detection covering all assets to ensure complete oversight",
    "Maintain comprehensive logs with minimum 12-month retention for complete forensic capability, addressing the logging gaps that hampered investigation",
    "Establish automated security incident response workflows to ensure consistent handling and timely reporting",
])
add_para(doc,
    'As MSIT Minister Yoo Sang-im emphasized, this breach serves as "a wake-up call for Korea\'s entire networked ecosystem" requiring organizations to "elevate information security to the very top of its management agenda" (MSIT, 2025). '
    'The incident demonstrates that, as noted in external analysis, "security on core infrastructure must be absolute and multi-layered" (Law and Ethics in Tech, 2025), requiring defense in depth rather than relying on single security mechanisms.'
)

# Conclusion
add_heading(doc, "Conclusion", 2)
add_para(doc,
    "The SK Telecom breach demonstrates catastrophic failures across multiple security principles, particularly Safe Defaults and Complete Mediation. "
    "The massive scale—affecting nearly 27 million subscribers—underscores the critical importance of proper security design in telecommunications infrastructure. "
    "By addressing the root causes identified in the official investigation and implementing security by design principles, organizations can build resilient security architectures that prevent similar breaches. "
    'The incident serves as a stark reminder that, as one analysis noted, "market dominance demands heightened accountability and customer care" (Law and Ethics in Tech, 2025), requiring robust security practices regardless of market position.'
)

# Works Cited (separate last page with hanging indent)
doc.add_page_break()
wc_title = add_heading(doc, "Works Cited", 2)
wc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

citations = [
    'MSIT. "MSIT Releases Final Investigation Results on SK Telecom Data Breach." Ministry of Science and ICT, 2025, https://www.msit.go.kr/eng/bbs/view.do;jsessionid=A2aV3fQR4zqYv-G8cJpkDgnrgrACDgREHvXAqG5l.AP_msit_2?sCode=eng&nttSeqNo=1139&bbsSeqNo=42&mId=4&mPid=2',
    'Law and Ethics in Tech. "Lessons from the SKT Hack: A Wake-Up Call for Digital Security and Corporate Responsibility." Medium, 2025, https://lawnethicsintech.medium.com/lessons-from-the-skt-hack-a-wake-up-call-for-digital-security-and-corporate-responsibility-9aa90b3ced41',
    'MITRE Corporation. "CWE-256: Plaintext Storage of Credentials." Common Weakness Enumeration, https://cwe.mitre.org/data/definitions/256.html',
    'MITRE Corporation. "CWE-312: Cleartext Storage of Sensitive Information." Common Weakness Enumeration, https://cwe.mitre.org/data/definitions/312.html',
    'MITRE Corporation. "CWE-798: Use of Hard-coded Credentials." Common Weakness Enumeration, https://cwe.mitre.org/data/definitions/798.html',
]

for c in citations:
    p = doc.add_paragraph(c, style="Bibliography")
    # Ensure single spacing and proper margins preserved
    p.paragraph_format.line_spacing = 1.0

# Save
output_path = "[Working] Report_2025952644_luke.docx"
doc.save(output_path)
output_path
